from django.shortcuts import render, redirect
import pandas as pd
from django.db.models import Q
from django.core.mail import send_mail
from django.conf import settings
import os
import socket
import docx
from django.conf import settings
from .models import Soha

socket.getaddrinfo("localhost", 8080)
# import textdistance

from django.http import HttpResponse, HttpResponseRedirect, JsonResponse

import re
from django.shortcuts import render

# Create your views here.
def mail(request):
    subject = "Thank you for registering to our site"
    message = " it  means a world to us "
    email_from = settings.EMAIL_HOST_USER
    recipient_list = [
        "mirzobek.tuliyev@gmail.com",
    ]
    send_mail(subject, message, email_from, recipient_list)
    return redirect("redirect to a new page")


def semanticTahlil(request):
    return render(request, "semanticTahlil.html")


def index(request):
    return render(request, "index.html")


def statistic(request):
    return render(request, "statistic.html")


def about(request):
    return render(request, "about.html")


def programService(request):
    return render(request, "program.html")


def electronicService(request):
    return render(request, "electronic.html")


def modelingService(request):
    return render(request, "modeling.html")


def contact(request):
    return render(request, "contact.html")


def service(request):
    return render(request, "service.html")


def education(request):
    return render(request, "education.html")


def hardwareProducts(request):
    return render(request, "hardwareProducts.html")


def simulationProducts(request):
    return render(request, "simulationProducts.html")


def tutorialProducts(request):
    return render(request, "tutorialProducts.html")


def import_L22(request):
    # # return HttpResponse("Already done!")

    # file = r"d:\Projects\Unicon-Lugat\slovari\L 21 Русско-узбекский толковый словарь терминов по телекоммуникационным услугам.docx"
    file = str(settings.MEDIA_ROOT) + "L 22 Русско-узбекский толковый словарь терминов по вычислительной технике.docx"

    doc = docx.Document(file)

    tables = doc.tables
    # q = 0
    for table in tables:
        for row in table.rows:
            # jadvaldagi sarlavha qatorlarni olmaslik uchun
            if len(row.cells[1].text) < 3:
                continue
            try:
                list = row.cells[0].text.splitlines()
                ru = list[0]
                uzl = ""
                uzk = ""
                en = ""
                for l in list[1:]:
                    if l.startswith("uz -"):
                        uzl = uzl + l[4:].strip()
                    elif l.startswith("en -"):
                        en = en + l[4:].strip()
                    else:
                        uzk = uzk + l.strip()

                # tarifini olish
                desc = row.cells[1].text.split("\n\n")
                desc_ru = desc[0]
                desc_uzl = desc[1]
                desc_uzk = desc[2]
                soha = Soha.objects.get(id=1)
                soha.soha_to_lugat.create(wordLatin=uzl, wordKiril=uzk, wordRus=ru, wordEnglish=en, commentLatin=desc_uzl, commentKiril=desc_uzk, commentRus=desc_ru)
            except:
                pass
    print("galdi")
    import_l1("e")
    import_l2("e")
    import_l4("e")
    import_l5("e")
    import_l6("e")
    import_l8("e")
    import_l9("e")
    import_l10("e")
    import_l11("e")
    import_l12("e")
    import_l13("e")
    import_l14("e")
    import_l15("e")
    import_l16("e")
    import_l17("e")
    import_l18("e")
    import_l19("e")
    import_l20("e")
    import_l21("e")
    import_l23("e")
    import_l24("e")
    import_l25("e")
    return HttpResponse("Done!")


def import_l1(request):
    try:
        file = str(settings.MEDIA_ROOT) + "L 1 Краткий русско-узбекский словарь по телекоммуникациям.docx"
        doc = docx.Document(docx=file)
        frond_word = ''
        r = 0
        for para in doc.paragraphs:
            if " — " in para.text:
                r += 1
                words = para.text.split(" — ")
                if " " in words[0][0]:
                    word_ru = (frond_word + ' ' + words[0].strip()).replace('  ', ' ')
                    word_uz_kr = words[1].strip().replace('  ', ' ')
                else:
                    word_ru = words[0].strip().replace('  ', ' ')
                    word_uz_kr = words[1].strip().replace('  ', ' ')
                    frond_word = word_ru.replace('  ', ' ').replace(' (', '(')
                try:
                    soha = Soha.objects.get(id=2)
                    soha.soha_to_lugat.create(wordKiril=word_uz_kr, wordRus=word_ru)
                except:
                    pass
    except Exception as ex:
        print(ex)
    # return HttpResponse("Already done!")


def import_l2(request):
    file = str(settings.MEDIA_ROOT) + "L 2 Словарь терминов по метрологии в сфере связи и информатизации.docx"
    doc = docx.Document(docx=file)
    tables = doc.tables
    for table in tables:
        for row in table.rows:
            if "Атама" in row.cells[1].text:
                print("galdi getti")
                continue
            try:
                # tarifini olish
                words = row.cells[1].text.strip()
                words_ru_index = words.find("ru - ")
                words_en_index = words.find("en - ")

                word_uz_kr = words[:words_ru_index].strip().replace('\n', ' ').replace('  ', ' ')
                word_ru = words[words_ru_index:words_en_index].strip().replace('\n', ' ').replace('  ', ' ')[5:]
                word_en = words[words_en_index:].strip().replace('\n', ' ').replace('  ', ' ')[5:]
                desc = row.cells[2].text.strip()
                soha = Soha.objects.get(id=3)
                soha.soha_to_lugat.create(wordKiril=word_uz_kr, wordRus=word_ru, wordEnglish=word_en, commentKiril=desc)
            except:
                pass
    # return HttpResponse("Already done!")


# def import_l3(request):
#     file = str(settings.MEDIA_ROOT) + "L 3 Русско-узбекский толковый словарь.docx"
#     doc = docx.Document(docx=file)
#
#     dont = ['А', 'Б', 'В', 'Г', 'Д', 'Е', 'Ё', 'Ж', 'З', 'И', 'Й', 'К', 'Л', 'М', 'Н', 'О', 'П', 'Р', 'С', 'Т', 'У',
#             'Ф', 'Х', 'Ц', 'Ч', 'Ш', 'Щ', 'Ъ', 'Ы', 'Ь', 'Э', 'Ю', 'Я']
#
#     tables = doc.tables
#     for table in tables:
#         for row in table.rows:
#             if len(row.cells) == 2:
#                 terms = str(row.cells[0].text.strip())
#                 texts = str(row.cells[1].text.strip())
#                 if (" / " in terms or terms == terms.upper() and len(row.cells) == 2 and terms not in dont and terms
#                         or "ТВт  ТВт" in terms):
#                     pass
#                 elif "Ы " in terms and "E " in terms:
#
#                     terms = terms.replace("\n", " ").replace("  ", " ")
#                     term_ru = terms[:terms.index('Ы ')]
#                     term_en = terms[terms.index('E '):][2:]
#                     term_uz = terms[terms.index('Ы '):terms.index('E ')][2:]


def import_l4(request):
    file = str(settings.MEDIA_ROOT) + "L 4 Электромагнитная-совместимость-радиотехнических-средств.-Термины-и-определения.docx"
    doc = docx.Document(docx=file)

    tables = doc.tables
    for table in tables:
        for row in table.rows:
            terms = str(row.cells[1].text.strip())
            if "ru -" in terms or "en -" in terms:  # ('ru – ' in terms or 'ru -' in terms) and ("en - " in terms or 'en – ' in terms or 'en  - ' in terms)
                texts = str(row.cells[-1].text.strip())
                if len([item for item in texts.split(".\n\n") if item]) == 2:
                    text_uz = texts.split(".\n\n")[0].replace('\n', '').strip()
                    text_ru = texts.split(".\n\n")[1].replace('\n', '').strip()
                else:  # if len([item for item in texts.split(". \n\n") if item]) == 2:
                    text_uz = texts.split(". \n\n")[0].replace('\n', '').strip()
                    text_ru = texts.split(". \n\n")[1].replace('\n', '').strip()

                term_uz = terms[:terms.index('ru -')].strip()
                term_en = terms[terms.index('en -'):][3:].strip()
                term_ru = terms[terms.index('ru -'):terms.index('en -')][3:].strip()
                try:
                    soha = Soha.objects.get(id=5)
                    soha.soha_to_lugat.create(wordKiril=term_uz, wordRus=term_ru, wordEnglish=term_en,
                                              commentKiril=text_uz, commentRus=text_ru)
                except:
                    pass
    # return HttpResponse("Already done!")


def import_l5(request):
    file = str(settings.MEDIA_ROOT) + "L 5 Словарь по электронной технике и радиоэлектронике.docx"
    doc = docx.Document(docx=file)

    tables = doc.tables
    for table in tables:
        for row in table.rows:
            terms = str(row.cells[0].text.strip())
            if terms == str(row.cells[1].text.strip()):
                pass
            elif 'en - ' in terms and 'ru - ' in terms:
                texts = str(row.cells[-1].text.strip())
                text_uz = ''
                text_ru = ''
                if len([item for item in texts.split(".\n\n") if item]) == 2:
                    text_uz = [item for item in texts.split(".\n\n") if item][0]
                    text_ru = [item for item in texts.split(".\n\n") if item][1]
                elif len([item for item in texts.split(". \n\n") if item]) == 2:
                    text_uz = [item for item in texts.split(". \n\n") if item][0]
                    text_ru = [item for item in texts.split(". \n\n") if item][1]
                elif len([item for item in texts.split(".\n") if item]) == 2:
                    text_uz = [item for item in texts.split(".\n") if item][0]
                    text_ru = [item for item in texts.split(".\n") if item][1]
                elif len([item for item in texts.split("\n\n") if item]) == 2:
                    text_uz = [item for item in texts.split("\n\n") if item][0]
                    text_ru = [item for item in texts.split("\n\n") if item][1]
                else:
                    pass

                term_uz = terms[:terms.index('ru - ')].replace('\n', '').strip()
                term_en = terms[terms.index('en - '):].replace('\n', '').strip()
                term_ru = terms[terms.index('ru - '):terms.index('en - ')].replace('\n', '').strip()

                try:
                    soha = Soha.objects.get(id=6)
                    soha.soha_to_lugat.create(wordKiril=term_uz, wordRus=term_ru, wordEnglish=term_en,
                                              commentKiril=text_uz, commentRus=text_ru)
                except:
                    pass
    # return HttpResponse("Already done!")


def import_l6(request):
    file = str(settings.MEDIA_ROOT) + "L 6 Русско-узбекский толковый словарь терминов по телевидению.docx"
    doc = docx.Document(docx=file)

    tables = doc.tables
    needs_row = []
    needs_desc = []
    for table in tables:
        rows = table.rows
        for row in rows:
            ter = row.cells[0].text.strip()
            if "ru - " in ter and "en - " in ter:
                needs_row.append(ter)
                needs_desc.append(row.cells[-1].text)

    check = False

    for term, desc in zip(needs_row, needs_desc):
        term_uz = term[:term.index('ru - ')].replace('\n', '').strip()
        term_en = term[term.index('en - '):].replace('\n', '').strip()[5:].strip()
        term_ru = term[term.index('ru - '):term.index('en - ')].replace('\n', '').strip()[5:].strip()

        if len([item for item in desc.split(".\n\n") if item]) == 2:
            text_uz = [item for item in desc.split(".\n\n") if item][0].strip().replace('\n', '')
            text_ru = [item for item in desc.split(".\n\n") if item][1].strip().replace('\n', '')

        elif len([item for item in desc.split(". \n\n") if item]) == 2:
            text_uz = [item for item in desc.split(". \n\n") if item][0].strip().replace('\n', '')
            text_ru = [item for item in desc.split(". \n\n") if item][1].strip().replace('\n', '')

        else:  # len([item for item in desc.split(".\n") if item]) == 2
            text_uz = [item for item in desc.split(".\n") if item][0].strip().replace('\n', '')
            text_ru = [item for item in desc.split(".\n") if item][1].strip().replace('\n', '')
        try:
            soha = Soha.objects.get(id=7)
            soha.soha_to_lugat.create(wordKiril=term_uz, wordRus=term_en, wordEnglish=term_ru,
                                      commentKiril=text_uz, commentRus=text_ru)
        except:
            pass

    # return HttpResponse("Already done!")


# def import_l7(request):
#     file = str(settings.MEDIA_ROOT) + "L 7 Словарь сокращений по телекоммуникациям.docx"
#     doc = docx.Document(docx=file)
#
#     row = 0
#     row2 = 0
#     tables = doc.tables
#     for i in tables:
#         row2 += len(i.rows)
#         for r in i.rows:
#             if r.cells[1].text == r.cells[2].text:
#                 row += 1
#             elif r.cells[0] == "Қисқартмалар":
#                 pass
#             else:
#                 shorts = r.cells[0].text.strip()
#                 uz = r.cells[3].text.strip()
#                 en = r.cells[1].text.strip()
#                 ru = r.cells[2].text.strip()
#                 try:
#                     soha = Soha.objects.get(id=5)
#                     soha.soha_to_lugat.create(wordKiril=term_uz, wordRus=term_ru, wordEnglish=term_en,
#                                               commentKiril=text_uz, commentRus=text_ru)
#                 except:
#                     pass
#
#     # return HttpResponse("Already done!")


def import_l8(request):
    file = str(settings.MEDIA_ROOT) + "L 8 Русско узбекский толковый словарь терминов по системам мобильной связи.docx"
    doc = docx.Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            if row.cells[1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[1] == ''):
                pass
            elif "uz -" in row.cells[0].text or "ru -" in row.cells[0].text:
                tex = str(row.cells[0].text)
                desc = str(row.cells[1].text)
                if 'uz -' in tex:
                    term_uz = tex[tex.index('uz -'):tex.index('en -')][4:].strip()
                    term_en = tex[tex.index('en -'):][5:].strip()
                    term_ru = tex[:tex.index('uz -')].strip()
                else:
                    term_uz = tex[tex.index('ru -'):tex.index('en -')][4:].strip()
                    term_en = tex[tex.index('en -'):][5:].strip()
                    term_ru = tex[:tex.index('ru -')].strip()

                if len(desc.split('. \n\n')) == 2:
                    desc_uz = desc.split('. \n\n')[1]
                    desc_ru = desc.split('. \n\n')[0]
                elif len(desc.split('.\n\n')) == 2:
                    desc_uz = desc.split('.\n\n')[1]
                    desc_ru = desc.split('.\n\n')[0]
                try:
                    soha = Soha.objects.get(id=9)
                    soha.soha_to_lugat.create(wordKiril=term_uz, wordRus=term_ru, wordEnglish=term_en,
                                              commentKiril=desc_uz, commentRus=desc_ru)
                except:
                    pass
            else:
                pass

    # return HttpResponse("Already done!")


def import_l9(request):
    file = str(settings.MEDIA_ROOT) + "L 9 Русско узбекский толковый словарь терминов по линиям связи и системам передачи.docx"
    doc = docx.Document(docx=file)

    tables = doc.tables
    for i in tables:
        for row in i.rows:
            if row.cells[1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[1] == ''):
                pass
            elif "uz - " in row.cells[0].text or "en - " in row.cells[0].text:
                tex = str(row.cells[0].text)
                desc = str(row.cells[1].text)

                term_uz = tex[tex.index('uz - '):tex.index('en - ')][5:].strip()
                term_en = tex[tex.index('en - '):][5:].strip()
                term_ru = tex[:tex.index('uz - ')].strip()

                if len(desc.split('. \n\n')) == 2:
                    desc_uz = desc.split('. \n\n')[1]
                    desc_ru = desc.split('. \n\n')[0]
                elif len(desc.split('.\n\n')) == 2:
                    desc_uz = desc.split('.\n\n')[1]
                    desc_ru = desc.split('.\n\n')[0]
                try:
                    soha = Soha.objects.get(id=10)
                    soha.soha_to_lugat.create(wordKiril=term_uz, wordRus=term_ru, wordEnglish=term_en,
                                              commentKiril=desc_uz, commentRus=desc_ru)
                except:
                    pass
    # return HttpResponse("Already done!")


def import_l10(request):
    file = str(settings.MEDIA_ROOT) + "L 10 Русско-узбекский толковый словарь терминов по электропитанию телекоммуникационных устройств.docx"
    doc = docx.Document(docx=file)

    tables = doc.tables
    for i in tables:
        for row in i.rows:
            if row.cells[1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[1] == ''):
                pass
            elif "uz - " in row.cells[0].text and "en - " in row.cells[0].text:
                tex = str(row.cells[0].text)
                desc = str(row.cells[1].text)

                term_uz = tex[tex.index('uz - '):tex.index('en - ')][5:].strip()
                term_en = tex[tex.index('en - '):][5:].strip()
                term_ru = tex[:tex.index('uz - ')].strip()

                if len(desc.split('. \n\n')) == 2:
                    desc_uz = desc.split('. \n\n')[1]
                    desc_ru = desc.split('. \n\n')[0]
                elif len(desc.split('.\n\n')) == 2:
                    desc_uz = desc.split('.\n\n')[1]
                    desc_ru = desc.split('.\n\n')[0]
                try:
                    soha = Soha.objects.get(id=11)
                    soha.soha_to_lugat.create(wordKiril=term_uz, wordRus=term_ru, wordEnglish=term_en,
                                              commentKiril=desc_uz, commentRus=desc_ru)
                except:
                    pass
    # return HttpResponse("Already done!")


def import_l11(request):
    file = str(settings.MEDIA_ROOT) + "L 11 Англо русско узбекский толковый словарь. Информационная технология. Операционные системы.docx"
    doc = docx.Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass
            elif "ru - " in row.cells[0].text and "uz - " in row.cells[0].text:
                tex = str(row.cells[0].text)
                desc = str(row.cells[-1].text)

                term_uz = tex[tex.index('uz - '):tex.index('ru - ')][5:].strip()
                term_en = tex[tex.index('ru - '):][5:].strip()
                term_ru = tex[:tex.index('uz - ')].strip()

                if len(desc.split('. \n\n')) == 2:
                    desc_uz = desc.split('. \n\n')[1]
                    desc_ru = desc.split('. \n\n')[0]
                elif len(desc.split('.\n\n')) == 2:
                    desc_uz = desc.split('.\n\n')[1]
                    desc_ru = desc.split('.\n\n')[0]

                try:
                    soha = Soha.objects.get(id=12)
                    soha.soha_to_lugat.create(wordKiril=term_uz, wordRus=term_ru, wordEnglish=term_en,
                                              commentKiril=desc_uz, commentRus=desc_ru)
                except:
                    pass

    # return HttpResponse("Already done!")


def import_l12(request):
    file = str(settings.MEDIA_ROOT) + "L 12 Словарь по информационной безопасности 1 издание.docx"
    doc = docx.Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en - " in row.cells[0].text and "uz - " in row.cells[0].text:
                tex = str(row.cells[0].text)
                desc = str(row.cells[-1].text)

                term_uz = tex[tex.index('uz - '):tex.index('en - ')][5:].strip()
                term_en = tex[tex.index('en - '):][5:].strip()
                term_ru = tex[:tex.index('uz - ')].strip()

                if len(desc.split('. \n\n')) == 2:
                    desc_uz = desc.split('. \n\n')[1]
                    desc_ru = desc.split('. \n\n')[0]
                elif len(desc.split('.\n\n')) == 2:
                    desc_uz = desc.split('.\n\n')[1]
                    desc_ru = desc.split('.\n\n')[0]
                try:
                    soha = Soha.objects.get(id=13)
                    soha.soha_to_lugat.create(wordKiril=term_uz, wordRus=term_ru, wordEnglish=term_en,
                                              commentKiril=desc_uz, commentRus=desc_ru)
                except:
                    pass
    # return HttpResponse("Already done!")


def import_l13(request):
    file = str(settings.MEDIA_ROOT) + "L 13 Русско узбекский толковый словарь терминов по радиотехнике.docx"
    doc = docx.Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en - " in row.cells[0].text and "uz - " in row.cells[0].text:
                tex = str(row.cells[0].text)
                desc = str(row.cells[-1].text)

                term_uz = tex[tex.index('uz - '):tex.index('en - ')][5:].strip()
                term_en = tex[tex.index('en - '):][5:].strip()
                term_ru = tex[:tex.index('uz - ')].strip()

                if len(desc.split('. \n\n')) == 2:
                    desc_uz = desc.split('. \n\n')[1]
                    desc_ru = desc.split('. \n\n')[0]
                elif len(desc.split('.\n\n')) == 2:
                    desc_uz = desc.split('.\n\n')[1]
                    desc_ru = desc.split('.\n\n')[0]
                try:
                    soha = Soha.objects.get(id=14)
                    soha.soha_to_lugat.create(wordKiril=term_uz, wordRus=term_ru, wordEnglish=term_en,
                                              commentKiril=desc_uz, commentRus=desc_ru)
                except:
                    pass
    # return HttpResponse("Already done!")


def import_l14(request):
    file = str(settings.MEDIA_ROOT) + "L 14 Русско узбекский толковый словарь терминов по коммутационным системам.docx"
    doc = docx.Document(docx=file)

    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en - " in row.cells[0].text and "uz - " in row.cells[0].text:

                term_uz = tex[tex.index('uz - '):tex.index('en - ')][5:].strip()
                term_en = tex[tex.index('en - '):][5:].strip()
                term_ru = tex[:tex.index('uz - ')].strip()

                if len(desc.split('. \n\n')) == 2:
                    desc_uz = desc.split('. \n\n')[1]
                    desc_ru = desc.split('. \n\n')[0]
                elif len(desc.split('.\n\n')) == 2:
                    desc_uz = desc.split('.\n\n')[1]
                    desc_ru = desc.split('.\n\n')[0]
                try:
                    soha = Soha.objects.get(id=15)
                    soha.soha_to_lugat.create(wordKiril=term_uz, wordRus=term_ru, wordEnglish=term_en,
                                              commentKiril=desc_uz, commentRus=desc_ru)
                except:
                    pass
    # return HttpResponse("Already done!")


def import_l15(request):
    file = str(settings.MEDIA_ROOT) + "L 15 Русско узбекский толковый словарь терминов по системам беспроводного доступа.docx"
    doc = docx.Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en - " in row.cells[0].text and "uz - " in row.cells[0].text:

                term_uz = tex[tex.index('uz - '):tex.index('en - ')][5:].strip()
                term_en = tex[tex.index('en - '):][5:].strip()
                term_ru = tex[:tex.index('uz - ')].strip()

                if len(desc.split('. \n\n')) == 2:
                    desc_uz = desc.split('. \n\n')[1]
                    desc_ru = desc.split('. \n\n')[0]
                elif len(desc.split('.\n\n')) == 2:
                    desc_uz = desc.split('.\n\n')[1]
                    desc_ru = desc.split('.\n\n')[0]
                try:
                    soha = Soha.objects.get(id=16)
                    soha.soha_to_lugat.create(wordKiril=term_uz, wordRus=term_ru, wordEnglish=term_en,
                                              commentKiril=desc_uz, commentRus=desc_ru)
                except:
                    pass
    # return HttpResponse("Already done!")


def import_l16(request):
    file = str(settings.MEDIA_ROOT) + "L 16 Русско узбекский толковый словарь терминов по измерениям в телекоммуникации.docx"
    doc = docx.Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass
            elif "en - " in row.cells[0].text and "uz - " in row.cells[0].text:

                term_uz = tex[tex.index('uz - '):tex.index('en - ')][5:].strip()
                term_en = tex[tex.index('en - '):][5:].strip()
                term_ru = tex[:tex.index('uz - ')].strip()

                if len(desc.split('. \n\n')) == 2:
                    desc_uz = desc.split('. \n\n')[1]
                    desc_ru = desc.split('. \n\n')[0]
                elif len(desc.split('.\n\n')) == 2:
                    desc_uz = desc.split('.\n\n')[1]
                    desc_ru = desc.split('.\n\n')[0]
                try:
                    soha = Soha.objects.get(id=17)
                    soha.soha_to_lugat.create(wordKiril=term_uz, wordRus=term_ru, wordEnglish=term_en,
                                              commentKiril=desc_uz, commentRus=desc_ru)
                except:
                    pass
    # return HttpResponse("Already done!")


def import_l17(request):
    file = str(settings.MEDIA_ROOT) + "L 17 Русско узбекский толковый словарь терминов по радиочастотному спектру_ радиоэлектронным средств.docx"
    doc = docx.Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]
                    try:
                        soha = Soha.objects.get(id=18)
                        soha.soha_to_lugat.create(wordKiril=term_uz, wordRus=term_ru, wordEnglish=term_en,
                                                  commentKiril=desc_uz, commentRus=desc_ru)
                    except:
                        pass
    # return HttpResponse("Already done!")


def import_l18(request):
    file = str(settings.MEDIA_ROOT) + "L 18 Русско-узбекский толковый словарь терминов по сетям передачи данных.docx"
    doc = docx.Document(docx=file)

    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]
                    try:
                        soha = Soha.objects.get(id=19)
                        soha.soha_to_lugat.create(wordKiril=term_uz, wordRus=term_ru, wordEnglish=term_en,
                                                  commentLatin=desc_uz, commentKiril=desc_uz, commentRus=desc_ru)
                    except:
                        pass
    # return HttpResponse("Already done!")


def import_l19(request):
    file = str(settings.MEDIA_ROOT) + "L 19 Русско-узбекский толковый словарь терминов по программированию.docx"
    doc = docx.Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]
                    try:
                        soha = Soha.objects.get(id=20)
                        soha.soha_to_lugat.create(wordKiril=term_uz, wordRus=term_ru, wordEnglish=term_en,
                                                  commentLatin=desc_uz, commentKiril=desc_uz_kr, commentRus=desc_ru)
                    except:
                        pass
    # return HttpResponse("Already done!")


def import_l20(request):
    file = str(settings.MEDIA_ROOT) + "L 20 Русско узбекский толковый словарь терминов по спутниковой связи.docx"
    doc = docx.Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass
            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]
                    try:
                        soha = Soha.objects.get(id=21)
                        soha.soha_to_lugat.create(wordKiril=term_uz, wordRus=term_ru, wordEnglish=term_en,
                                                  commentLatin=desc_uz, commentKiril=desc_uz_kr, commentRus=desc_ru)
                    except:
                        pass
    # return HttpResponse("Already done!")


def import_l21(request):
    file = str(settings.MEDIA_ROOT) + "L 21 Русско-узбекский толковый словарь терминов по телекоммуникационным услугам.docx"
    doc = docx.Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]
                    try:
                        soha = Soha.objects.get(id=22)
                        soha.soha_to_lugat.create(wordKiril=term_uz, wordRus=term_ru, wordEnglish=term_en,
                                                  commentLatin=desc_uz, commentKiril=desc_uz_kr, commentRus=desc_ru)
                    except:
                        pass
    # return HttpResponse("Already done!")


def import_l23(request):
    file = str(settings.MEDIA_ROOT) + "L 23 Русско-узбекский толковый словарь терминов по электронному документообороту.docx"
    doc = docx.Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]
                    try:
                        soha = Soha.objects.get(id=23)
                        soha.soha_to_lugat.create(wordKiril=term_uz, wordRus=term_ru, wordEnglish=term_en,
                                                  commentLatin=desc_uz, commentKiril=desc_uz_kr, commentRus=desc_ru)
                    except:
                        pass
    # return HttpResponse("Already done!")


def import_l24(request):
    file = str(settings.MEDIA_ROOT) + "L 24 Русско-узбекский толковый словарь терминов по оптоэлектронике.docx"
    doc = docx.Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]
                    try:
                        soha = Soha.objects.get(id=24)
                        soha.soha_to_lugat.create(wordKiril=term_uz, wordRus=term_ru, wordEnglish=term_en,
                                                  commentLatin=desc_uz, commentKiril=desc_uz_kr, commentRus=desc_ru)
                    except:
                        pass
    # return HttpResponse("Already done!")


def import_l25(request):
    file = str(settings.MEDIA_ROOT) + "L 25 Русско-узбекский толковый словарь терминов по радиорелейным системам.docx"
    doc = docx.Document(docx=file)
    tables = doc.tables
    for i in tables:
        for row in i.rows:
            tex = str(row.cells[0].text)
            desc = str(row.cells[-1].text)
            if row.cells[-1].text == row.cells[0].text or (row.cells[0] == '' and row.cells[-1] == ''):
                pass

            elif "en -" in row.cells[0].text and "uz -" in row.cells[0].text:
                term_uz = tex[tex.index('uz -'):tex.index('en -')][5:].strip()
                term_en = tex[tex.index('en -'):][5:].strip()
                term_ru = tex[:tex.index('uz -')].strip()
                descs = [element for element in desc.split('\n\n') if element]
                if len(descs) == 3:
                    desc_uz = desc.split('\n\n')[1]
                    desc_uz_kr = desc.split('\n\n')[2]
                    desc_ru = desc.split('\n\n')[0]
                    try:
                        soha = Soha.objects.get(id=25)
                        soha.soha_to_lugat.create(wordKiril=term_uz, wordRus=term_ru, wordEnglish=term_en,
                                                  commentLatin=desc_uz, commentKiril=desc_uz_kr, commentRus=desc_ru)
                    except:
                        pass
    # return HttpResponse("Already done!")


def show_sohalar(request):
    from .models import Soha

    sohalar = Soha.objects.all()
    return render(request, "show_soha.html", context={"sohalar": sohalar})


def show_lugat(request):
    from .models import Lugat

    # sohalar = Soha.object.filter(id=request.GET["s"])
    lugat = Lugat.objects.select_related("soha").filter(soha=request.GET.get("s"))
    return render(request, "show_lugat.html", context={"lugat": lugat})
